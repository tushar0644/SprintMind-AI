import io
import re
from abc import ABC, abstractmethod
from typing import Dict, List, Optional
from pydantic import BaseModel
from pypdf import PdfReader
import docx

from app.documents.schemas import DocumentParserResponse, DocumentParserMetadata


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, file_bytes: bytes, filename: str) -> DocumentParserResponse:
        """
        Extracts clean plain text and returns a unified document response with metadata.
        """
        pass

    def detect_language(self, text: str) -> str:
        """
        Detects language using a fast, lightweight stopword frequency heuristic.
        Supports English, Spanish, German, French. Defaults to 'en'.
        """
        if not text or not text.strip():
            return "en"

        text_lower = text.lower()
        words = re.findall(r"\b\w+\b", text_lower)
        if not words:
            return "en"

        # Check top stopwords for each language
        stopwords: Dict[str, List[str]] = {
            "en": ["the", "be", "to", "of", "and", "a", "in", "that", "have", "i", "it", "for", "not", "on", "with", "he", "as", "you", "do", "at"],
            "es": ["el", "la", "los", "las", "un", "una", "de", "en", "que", "y", "o", "no", "se", "con", "por", "para", "como", "su", "al", "del"],
            "de": ["der", "die", "das", "und", "ist", "in", "zu", "den", "von", "mit", "ein", "eine", "sind", "es", "an", "auf", "für", "dem", "im"],
            "fr": ["le", "la", "les", "un", "une", "des", "de", "et", "en", "que", "est", "dans", "pour", "par", "qui", "ce", "sur", "plus", "une", "pas"]
        }

        counts = {lang: 0 for lang in stopwords.keys()}
        # Count match occurrences
        for word in words[:200]:  # analyze first 200 words for speed and accuracy
            for lang, stop_list in stopwords.items():
                if word in stop_list:
                    counts[lang] += 1

        best_lang = max(counts, key=counts.get)
        # If no stopwords matched, fallback to English
        if counts[best_lang] == 0:
            return "en"
        return best_lang

    def calculate_stats(self, text: str) -> tuple[int, int]:
        """
        Returns (word_count, character_count) for the given text.
        """
        if not text:
            return 0, 0
        word_count = len(text.split())
        character_count = len(text)
        return word_count, character_count


class PDFParser(DocumentParser):
    def parse(self, file_bytes: bytes, filename: str) -> DocumentParserResponse:
        if not file_bytes:
            return DocumentParserResponse(
                content="",
                metadata=DocumentParserMetadata(
                    title=filename,
                    pages=0,
                    word_count=0,
                    character_count=0,
                    language="en"
                )
            )

        pdf_file = io.BytesIO(file_bytes)
        try:
            reader = PdfReader(pdf_file)
            pages_count = len(reader.pages)
            
            text_list = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_list.append(page_text)
            
            content = "\n".join(text_list)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

        # Retrieve title from PDF metadata or fallback
        title = None
        if reader.metadata and reader.metadata.title:
            title = reader.metadata.title.strip()
        
        if not title:
            # Fallback to first line of text
            first_line = content.split("\n")[0].strip() if content else ""
            title = first_line[:80] if (first_line and len(first_line) > 3) else filename

        word_count, char_count = self.calculate_stats(content)
        language = self.detect_language(content)

        return DocumentParserResponse(
            content=content,
            metadata=DocumentParserMetadata(
                title=title,
                pages=pages_count,
                word_count=word_count,
                character_count=char_count,
                language=language
            )
        )


class DocxParser(DocumentParser):
    def parse(self, file_bytes: bytes, filename: str) -> DocumentParserResponse:
        if not file_bytes:
            return DocumentParserResponse(
                content="",
                metadata=DocumentParserMetadata(
                    title=filename,
                    pages=0,
                    word_count=0,
                    character_count=0,
                    language="en"
                )
            )

        docx_file = io.BytesIO(file_bytes)
        try:
            doc = docx.Document(docx_file)
            
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            
            # Extract text from tables to be thorough
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            content = "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

        # Retrieve title from DOCX core properties or fallback
        title = None
        if doc.core_properties and doc.core_properties.title:
            title = doc.core_properties.title.strip()
        
        if not title:
            # Fallback to first paragraph text
            first_para = doc.paragraphs[0].text.strip() if doc.paragraphs else ""
            title = first_para[:80] if (first_para and len(first_para) > 3) else filename

        word_count, char_count = self.calculate_stats(content)
        language = self.detect_language(content)

        # Docx pages count fallback
        pages = doc.core_properties.pages if (doc.core_properties and doc.core_properties.pages) else 0
        if pages <= 0:
            pages = max(1, word_count // 400)

        return DocumentParserResponse(
            content=content,
            metadata=DocumentParserMetadata(
                title=title,
                pages=pages,
                word_count=word_count,
                character_count=char_count,
                language=language
            )
        )


class TxtParser(DocumentParser):
    def parse(self, file_bytes: bytes, filename: str) -> DocumentParserResponse:
        if not file_bytes:
            return DocumentParserResponse(
                content="",
                metadata=DocumentParserMetadata(
                    title=filename,
                    pages=1,
                    word_count=0,
                    character_count=0,
                    language="en"
                )
            )

        try:
            content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                content = file_bytes.decode("latin-1")
            except Exception as e:
                raise ValueError(f"Failed to decode TXT file: {str(e)}")

        word_count, char_count = self.calculate_stats(content)
        language = self.detect_language(content)

        # Fallback title calculation
        first_line = content.split("\n")[0].strip() if content else ""
        title = first_line[:80] if (first_line and len(first_line) > 3) else filename

        return DocumentParserResponse(
            content=content,
            metadata=DocumentParserMetadata(
                title=title,
                pages=1,
                word_count=word_count,
                character_count=char_count,
                language=language
            )
        )


class MarkdownParser(DocumentParser):
    def clean_markdown(self, md_text: str) -> str:
        """
        Removes Markdown specific styling elements to leave clean plain text.
        """
        if not md_text:
            return ""

        # Remove code blocks (```code```)
        text = re.sub(r"```[\s\S]*?```", "", md_text)
        
        # Remove HTML tags
        text = re.sub(r"<[^>]*>", "", text)
        
        # Replace image links ![alt](url) with just alt
        text = re.sub(r"!\[(.*?)\]\(.*?\)", r"\1", text)
        
        # Replace links [text](url) with just text
        text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
        
        # Remove bold and italics (*, _, **, __, ~~, `)
        text = re.sub(r"(\*\*|__|\*|_|~~|`)", "", text)
        
        # Remove headings syntax (# Heading)
        text = re.sub(r"^\s*#+\s+", "", text, flags=re.MULTILINE)
        
        # Remove horizontal rules (---, ***, ___)
        text = re.sub(r"^\s*[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
        
        # Remove blockquote formatting
        text = re.sub(r"^\s*>\s+", "", text, flags=re.MULTILINE)

        # Clean list markers at the beginning of lines
        text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)

        # Remove duplicate empty lines and clean whitespace
        text = "\n".join([line.strip() for line in text.split("\n") if line.strip()])
        return text

    def parse(self, file_bytes: bytes, filename: str) -> DocumentParserResponse:
        if not file_bytes:
            return DocumentParserResponse(
                content="",
                metadata=DocumentParserMetadata(
                    title=filename,
                    pages=1,
                    word_count=0,
                    character_count=0,
                    language="en"
                )
            )

        try:
            content_raw = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                content_raw = file_bytes.decode("latin-1")
            except Exception as e:
                raise ValueError(f"Failed to decode Markdown file: {str(e)}")

        content = self.clean_markdown(content_raw)
        word_count, char_count = self.calculate_stats(content)
        language = self.detect_language(content)

        # Extracted Title: Find first h1 style heading (# Header)
        title = filename
        h1_match = re.search(r"^\s*#\s+(.+)$", content_raw, re.MULTILINE)
        if h1_match:
            title = h1_match.group(1).strip()
        else:
            first_line = content.split("\n")[0].strip() if content else ""
            if first_line and len(first_line) > 3:
                title = first_line[:80]

        return DocumentParserResponse(
            content=content,
            metadata=DocumentParserMetadata(
                title=title,
                pages=1,
                word_count=word_count,
                character_count=char_count,
                language=language
            )
        )


class DocumentParserFactory:
    @staticmethod
    def get_parser(filename: str) -> DocumentParser:
        ext = filename.split(".")[-1].lower() if "." in filename else ""
        if ext == "pdf":
            return PDFParser()
        elif ext == "docx":
            return DocxParser()
        elif ext == "txt":
            return TxtParser()
        elif ext in ["md", "markdown"]:
            return MarkdownParser()
        else:
            raise ValueError(f"Unsupported file format: '.{ext}'")
